import requests
from bs4 import BeautifulSoup
import os
from xml.etree import ElementTree as ET
import csv
import time
import urllib2
from bs4 import BeautifulSoup          
from docx import Document
import lxml
import os
import docx
import time
#showtime = strftime("%y-%m-%d,gmtime())
print (time.strftime("%d/%m/%y"))
split1=time.strftime("%d/%m/%y")
split1=split1.split("/")
#print(split1[0])
#print(split1[1])
#print(split1[2])
#rootdir ="C:\Python27\Lib\python-docx-0.8.6"
os.makedirs('D:\\result',0755)
os.makedirs('D:\\result\\details_csv',0755)
os.makedirs('D:\\result\\summary_csv',0755)
os.makedirs('D:\\result\\images',0755)
os.makedirs('D:\\result\\docx',0755)

#Folder Creation 
url=[]
x=open(r"C:\Users\rsi142\Documents\My Received Files\DE.txt")
count=0
url=[]
for t in x:
    t=t.strip(' \n\t')
    url.append(t)
    



name=[""]*18
name[0]="home"
name[1]="about-us"
name[2]="contact-us"
name[3]="news"
name[4]="support"
name[5]="warranty"
name[6]="mate9"
name[7]="p10"
name[8]="nova"
name[9]="watch2"
name[10]="band2-pro"
name[11]="matebook"
name[12]="mediapad-m2-10"
name[13]="mobile-broadband"
name[14]="matebook-x"
name[15]="mate10-pro"
name[16]="mate10-pro-specs"
name[17]="matebook-x-specs"
print(url[0])
count=0
for x in url:
  
  response = requests.get("http://www.webpagetest.org/runtest.php?url="+x+"&runs=1&f=xml&k=A.f01eb720697f661682eea8925a0e0792&location=Frankfurt_Ruxit.3GFast&mobile=on&mdev=Nexus5X")
  time.sleep(6)

#print(response.content)
  print(response.content)
  tree = ET.fromstring(response.content)
#root = tree.getroot()
  root=tree.findall("./data/")
# for x in child:
  # print x.tag,x.text
#root = ET.fromstring(tree)
#child=root.findall("/data")
#for child in data:
#  print(child.tag,child.text)
  data1=[]
  for child in root:
    print(child.tag,child.text)
    if child.tag=="xmlUrl":
      data1.append(child.text)
    if child.tag=="testId":
      data1.append(child.text)
    if child.tag=="userUrl":
      data1.append(child.text)
    if child.tag=="summaryCSV":
      data1.append(child.text)
    if child.tag=="jsonUrl":
      data1.append(child.text)  
    if child.tag=="detailCSV":
      data1.append(child.text)  
#for x in data1:
#  print(x)
  fill=requests.get(data1[1])
#print(fill.content)
  tree2=ET.fromstring(fill.content)
  t=tree2.getchildren()
  print(t[1].text,t[1].tag)
  while t[1].text!="Ok":
    time.sleep(500)
    print(t[1].text)
    fill=requests.get(data1[1])
    tree2=ET.fromstring(fill.content)
    t=tree2.getchildren()
#root1=tree2.findall("./response")
# for r in root1:
  # print(child.text,child.tag) 

#print "hello "
  print t[1].text
  t=data1[0]

  print(t)
  mr=t

  print(mr)
  r=t.split("_")
  print(r[1],r[2])
  for t in range(1,4):
    print(data1[t])
    response=requests.get(data1[t])
  #print(response.content)
#response1=requests.get(data1[2])  
#with open(response1, "wb") as f:
#        f.write(response.content)
	#testid=child.text
  temp_file_name =name[count]+".csv"
  
  url = data1[3]
  print(url)
  print("checking summary data here")
  download = requests.get(url)
#write the content into the csv

  path="D:\\result\\summary_csv\\"
  filepath = os.path.join(path,temp_file_name)
 
  with open(filepath, 'w') as temp_file:
      temp_file.writelines(download.content)

  with open(filepath, 'rU') as temp_file:
      csv_reader = csv.reader(temp_file, dialect=csv.excel_tab)
    #for line in csv_reader:
     #   print line	
  temp_file_name = name[count]+".csv"
  url = data1[4]
  print(url)
  #print("security check")
  download = requests.get(url)
  path="D:\\result\\details_csv\\"
  filepath = os.path.join(path,temp_file_name)

  #path=path.join(temp_file_name)
  with open(filepath, 'w') as temp_file:
      temp_file.writelines(download.content)

  with open(filepath, 'rU') as temp_file:
      csv_reader = csv.reader(temp_file, dialect=csv.excel_tab)
    #for line in csv_reader:
     #   print line			


  url_pic="http://www.webpagetest.org/results/"+split1[2]+"/"+split1[1]+"/"+split1[0]+"/"+r[1]+"/"+r[2]+"/"
  print(url_pic)
  pic_list=["1_optimization.png","1_Cached_waterfall.png","1_Cached_connection.png","1_waterfall.png","1_connection.png"]	
  for x in pic_list:
    
		temp_url=url_pic+x
		#url of the pic where the image is present
		#print(temp_url)
		path='D:\\result\\images\\'
		
		#path of the image 
		response=requests.get(temp_url)
		t=name[count]
		r=t+x
		filepath=os.path.join(path,r)
		with open(filepath,"wb") as f:
		  f.write(response.content)
  	

#response=requests.get("http://10.150.18.159//results//16//11//15//FW//8R2//1_optimization.png")
#http://10.150.18.159//results//16//11//15//1_optimization.png
#image_name="1_optimization.png"
#bold = "\033[1m"
  r="1"
  #print(t)
#mr=str(mr)
  url1="http://www.webpagetest.org/result/"+mr+"/1/performance_optimization/"
  #print(url1)

  response=urllib2.urlopen(url1).read()
#print(response)
  soup=BeautifulSoup(response,'html.parser')
  heading_tags=soup.find_all("h3")
  data1=[]
#ta=soup.find("a",name="progressive_jpeg")
#for x in ta:
  #print(x.text)
  for x in heading_tags:
  #print x.text.encode("ascii", errors="ignore").decode("ascii")
    data1.append(x.text.encode("ascii", errors="ignore").decode("ascii"))
#soup1=BeautifulSoup(response)
  para_tags=soup.find_all("p") 
#for x in para_tags:
#print(tweet.encode("ascii", errors="ignore").decode("ascii"))
#for x in para_tags:
  #print x.text.encode("ascii", errors="ignore").decode("ascii")
#  print x.get("b",None)
#soup2=BeautifulSoup(response) 
  class1=soup("class") 
  t=soup.find_all("p", class_="indented1")
  ta=soup.find_all("p")
  for x in t:
    print(x.text.encode("ascii", errors="ignore").decode("ascii"))
#print(t)
  l=0
  print("whole data")
#f=open('output.docx','wb')
  
  lw=name[count]+".docx"
  #doc = docx.Document(os.path.join(rootdir,l))
  document = Document()
#p=document.add_paragraph('') 
  for x in t:
 # print("hello")
	  #print("indented")
	  #print(data1[l])
	  
		
	  p=document.add_paragraph('')
	  p.add_run(data1[l]).bold = True
	
	  document.add_paragraph(x.text.encode("ascii", errors="ignore").decode("ascii"))
	  l=l+1
	  
	  path='D:\\result\\docx'
	  
	  filepath=os.path.join(path,lw)
  document.save(filepath)
  count=count+1
  

 
  #f.write()
  #print x.text.encode("ascii", errors="ignore").decode("ascii")
  #f.write(x.text.encode("ascii", errors="ignore").decode("ascii"))
  
  #p=len(x.contents)
  #print(x[0])
  




